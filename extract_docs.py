#!/usr/bin/env python3
"""Extract content from Excel and Word documents"""
import os
from openpyxl import load_workbook
from docx import Document

def extract_excel(filepath):
    """Extract all sheets from Excel file"""
    print(f"\n{'='*60}")
    print(f"FILE: {os.path.basename(filepath)}")
    print('='*60)
    try:
        wb = load_workbook(filepath, data_only=True)
        for sheet_name in wb.sheetnames:
            print(f"\n--- Sheet: {sheet_name} ---")
            ws = wb[sheet_name]
            for row in ws.iter_rows(max_row=50, values_only=True):
                if any(cell is not None for cell in row):
                    print(" | ".join(str(cell) if cell else "" for cell in row))
    except Exception as e:
        print(f"Error: {e}")

def extract_docx(filepath):
    """Extract content from Word document"""
    print(f"\n{'='*60}")
    print(f"FILE: {os.path.basename(filepath)}")
    print('='*60)
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        for table in doc.tables:
            print("\n[TABLE]")
            for row in table.rows:
                print(" | ".join(cell.text for cell in row.cells))
    except Exception as e:
        print(f"Error: {e}")

# Extract all department input files
input_dir = "DOCS/ALL DEPT DATA INPUT CELLS-20260105T072145Z-3-001/ALL DEPT DATA INPUT CELLS"
for f in os.listdir(input_dir):
    if f.endswith('.xlsx'):
        extract_excel(os.path.join(input_dir, f))

# Extract sample output files
output_dir = "DOCS/SAMPLE DOCUMENTS FOR OUTPUT-20260105T072149Z-3-001/SAMPLE DOCUMENTS FOR OUTPUT"
for f in os.listdir(output_dir):
    fp = os.path.join(output_dir, f)
    if f.endswith('.xlsx'):
        extract_excel(fp)
    elif f.endswith('.docx'):
        extract_docx(fp)

# Extract reading material docx files
reading_dir = "DOCS/READING MATERIAL-20260105T072148Z-3-001/READING MATERIAL"
for f in os.listdir(reading_dir):
    if f.endswith('.docx'):
        extract_docx(os.path.join(reading_dir, f))